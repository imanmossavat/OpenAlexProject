import logging
import re
from pathlib import Path
from typing import List, Optional

import docx

from .base import BaseExtractor
from ..models import PaperMetadata
from ..utils import parse_author_list, find_doi


class DocxExtractor(BaseExtractor):
    """Extract metadata from DOCX documents."""

    def __init__(self, logger: Optional[logging.Logger] = None) -> None:
        self._logger = logger or logging.getLogger(__name__)

    def extract(self, path: str) -> PaperMetadata:
        path_obj = Path(path)
        if not path_obj.exists():
            self._logger.error("DOCX not found: %s", path)
            return PaperMetadata()

        try:
            document = docx.Document(str(path_obj))
        except Exception as exc:
            self._logger.error("Unable to open DOCX %s: %s", path, exc)
            return PaperMetadata()

        paragraphs = self._collect_text(document)
        if not paragraphs:
            return PaperMetadata()

        title = self._extract_title(paragraphs)
        authors = self._extract_authors(paragraphs)
        abstract = self._extract_abstract(paragraphs)
        year = self._extract_year(paragraphs)
        doi = find_doi("\n".join(paragraphs))

        return PaperMetadata(
            title=title,
            authors=authors,
            abstract=abstract,
            year=year,
            doi=doi,
            venue=None,
        )

    def _collect_text(self, document) -> List[str]:
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        paragraphs.append(txt)

        return paragraphs

    def _extract_title(self, paragraphs: List[str]) -> Optional[str]:
        return paragraphs[0] if paragraphs else None

    def _extract_authors(self, paragraphs: List[str]) -> List[str]:
        if not paragraphs:
            return []

        author_patterns = [
            r"^authors?\s*[:\-–—]\s*(.+)$",
            r"^authors?\s+(.+)$",
            r"^author\(s\)\s*[:\-–—]\s*(.+)$",
        ]

        for p in paragraphs:
            for pattern in author_patterns:
                match = re.match(pattern, p, flags=re.IGNORECASE)
                if match:
                    return parse_author_list(match.group(1))

        for idx, p in enumerate(paragraphs):
            if p.lower() == "authors" and idx + 1 < len(paragraphs):
                return parse_author_list(paragraphs[idx + 1])

        if len(paragraphs) > 1:
            candidate = paragraphs[1]
            if re.search(r"[A-Za-z]+\s+[A-Za-z]+", candidate):
                return parse_author_list(candidate)

        return []

    def _extract_abstract(self, paragraphs: List[str]) -> Optional[str]:
        for i, p in enumerate(paragraphs):
            if p.lower().startswith("abstract"):
                if ":" in p:
                    inline = p.split(":", 1)[1].strip()
                    if inline:
                        return inline
                if i + 1 < len(paragraphs):
                    return paragraphs[i + 1]

        if len(paragraphs) > 2:
            candidate = paragraphs[2]
            if not re.match(r"^\d{4}$", candidate):
                return candidate

        return None

    def _extract_year(self, paragraphs: List[str]) -> Optional[str]:
        for p in paragraphs:
            if p.lower().startswith("year"):
                value = p.split(":", 1)[-1].strip()
                if re.match(r"^\d{4}$", value):
                    return value

        for idx, p in enumerate(paragraphs):
            if p.lower() == "year" and idx + 1 < len(paragraphs):
                value = paragraphs[idx + 1]
                if re.match(r"^\d{4}$", value):
                    return value

        for p in paragraphs:
            if re.match(r"^\d{4}$", p):
                return p

        return None
